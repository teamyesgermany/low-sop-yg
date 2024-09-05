import streamlit as st
import io
import os
import openai as ai
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
import spacy
import json
import random
import pickle
import asyncio
import aiohttp
import pyrebase
import urllib.parse
import warnings
from datetime import datetime, timedelta
from streamlit_cookies_manager import EncryptedCookieManager
import requests
from bs4 import BeautifulSoup
import sys
# import logging

#! Configure the logging functionnality 
# logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

#* Initialization of environment variables for storing API KEYS and the PICKLE FILE

PICKLE_FILE = 'scraped_data.pkl'
load_dotenv()  # take environment variables from .env.
ai.api_key = os.getenv("OPENAI_API_KEY")
PYREBASE_API_KEY = os.getenv("PYREBASE_API_KEY")
DIFFBOT_API_TOKEN = os.getenv("DIFFBOT_API_KEY")
DIFFBOT_API_URL = 'https://api.diffbot.com/v3/article'

#& Firebase authentication and configuration settings

firebaseConfig = {
    'apiKey': f"{PYREBASE_API_KEY}",
    'authDomain': "yg----sop---high.firebaseapp.com",
    'projectId': "yg----sop---high",
    'databaseURL':"https://yg----sop---high-default-rtdb.firebaseio.com/",
    'storageBucket': "yg----sop---high.appspot.com",
    'messagingSenderId': "203643126427",
    'appId':  "1:203643126427:web:54ce1ef165fc2b40993c3e",
    'measurementId': "G-2QQBJK5FG5"
}
# Firebase authentication
firebase = pyrebase.initialize_app(firebaseConfig)
auth = firebase.auth()


#^ Load the SpaCy English model + creation of english blank model for SPACY TECHNOLOGY

nlp = spacy.blank("en")  # Create a blank English model
nlp_e = spacy.load("en_core_web_sm")


#! URL path of your logo image
logo_path = 'images/yeslogo2.jpg'



#!!! the function that generates the SOP :
    

async def generate_sop_version1(template_text, res_text,programme,university):
    
    
    cgpa_score = await retrieve_cgpa_score(res_text)
    
    print("from the terminal : ", cgpa_score)
    
    if cgpa_score != 'unknown' and float(cgpa_score) >= 7: 
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
        messages=[
             {"role": "user", "content": "Generate a statement of purpose (SOP) based on the provided details, following a specific structure and style."},

            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Here is the resume {res_text} from which you have to take information about the student and please make sure that you do not add things from your own and stick to this {res_text} througout the sop"},
            {"role": "user" , "content": f"Another important that you should keep in mind while making the sop is that the sop is a letter of motivation to study in this {university} , so you have to generate the sop which should contian strong motivational points for choosing {programme}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"University name: {university}"},

            {"role": "user", "content": f" For the Introduction's paragraph, you must be using content from the {res_text} and follow the same structure and catchy words from from the introduction paragraph of this template {template_text}"},


            {"role": "user", "content": f"""Fot the Academic and experience Paragraph too you will be using content from the {res_text}: Discuss the candidate's academic and professional background :
            IF CGPA score is more than 7 and his Bachelor graduation date is PAST : Mention his CGPA score.
            Mention all his acadamic education scores from the {res_text}
             Mention all the acadamic education scores as they are prensent in {res_text}
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list.
            You will be mentioning things like his certificates and Ilets score if they are present in the {res_text}
            """},


            {"role": "user", "content": f""" For the paragraph about the programme use strong motivational points for choosing {programme} and add information from the programme into this paragraph, 
                Show a STRONG motivation by evoking a storytelling from your past that led you want to study this programme.
                If the past education of yours is related to the programme you are applying then you can consider them adding in this paragraph too
                You have to choose MINIMUM two out of these points :
                Story from a family experience
                Story from a social phenomenon and you want to bring solutions
                Story from previous internship
                Story from the studied modules in your Academic cursus
                Story from one of your extra-curricular activites
                You will also be adding What are the benifits you will be getting from this {programme} and the {university} 
                You must also select some modules proposed by the programme and describe their contents, then explain why you are HIGHLY motivated to study them.
                Mention the learning outcomes from this {programme}, which include technical knowledge and interpersonal skills.     
                How this course will benefit me professionally and for persona growth as well.
                Explain Why this programme can help you build your career and prepare you for your future too.
                Please write this paragraph in the first-person tone where I express my strong motivation to study in Germany. 
                Start each sentence with 'I' or include 'my'
                **Important:** Please do not invent any information not found in the `programme_response`. Base your content strictly on the provided information. If any required information is missing from the source, simply omit that point rather than fabricating details.

            """},

            {"role": "user", "content": f"""For the paragraph about the University, you need to craft a compelling and personalized statement about why you want to apply to {university}. Write in the first person, as if you are the applicant. Incorporate detailed information from into this paragraph, addressing the following points:

            1. **University Ranking and Source:**
            - Include the exact ranking of {university} and cite the source of this ranking.

            2. **Student Body and Facilities:**
            - Mention the number of students at {university}.
            - Describe the facilities, faculty, and campus location in detail.

            3. **Personal Connections:**
            - Note if you have any relatives or friends at {university} and how that influences your choice.

            4. **Research Centers and Professors:**
            - Highlight any notable research centers and renowned professors associated with the program you are interested in.

            5. **City Attractions:**
            - Explain why you want to study in the city where {university} is located. Include some cool spots or attractions in the city that interest you.

            6. **Local Job Market:**
            - Provide examples of companies in the city that hire graduates from the {programme}.
            Based on the above 6 points, add the following points in the paragraph:
            7. Academic benifits:
            - Mention the academic benifits of studying in this {university}

            8. Personal benifits:
            - Mention the personal benifits of studying in this {university}

            9. Professional benifits:
            -Mention the professional benifits of studying in this {university}

            Please write this paragraph in the first-person tone where I express my strong motivation to study in Germany. 
            Start each sentence with 'I' or include 'my'

            Make sure to weave all these elements into a cohesive and engaging paragraph that reflects your genuine interest in both the university and the city. Be detailed, specific, and personal in your response.

            **Important:** Please do not invent any information not found in the `university_response`. Base your content strictly on the provided information. If any required information is missing from the source, simply omit that point rather than fabricating details.

            """},
            {"role": "user", "content": f"""For the paragraph which is about your future career perspectives and aspirations post-study, you should mention how the programme of {programme}
             will give you the opportunity to work in the futur job you want to do. You should mention names of professional positions linked to the programme you want to work in.
             You can also add that you aim to work in Germany"""},


            {"role": "user", "content": "Last Paragraph: A brief conclusion summarizing why you are the ideal applicant and show again your interest."},
            {"role": "user", "content": f"""
            For the paragraph about Germany, you must mention the following points and add information from  the news into this paragraph :
            
            You strictly have to include at least 5-6 good points for this paragraph.
            Point 1. Explain why you chose to apply in Germany
            Point 2. Your Intention to stay in Germany because the field education comparative is far better to others destinations
            Point 3. Talk about the Good exposure, diversity and German culture 
            point 4. What are the reasons to choose Germany for masters/bachelors program?
            point 5. How studying in Germnay will help me grow professionally as well as personally
            point 6. Facts about Germany and its culture that led to the decision to study in Germany
            Please ensure each sentence transitions smoothly into the next, maintaining a logical flow throughout the document.
            Please write this paragraph in the first-person tone where I express my strong motivation to study in Germany. 
            Start each sentence with 'I' or include 'my'
            please can you write each sentence like "I want to …." instead of "Germany's education system is highly regarded for its quality and affordability" and please write the whole paragraph like this only.
            **Important:** Please do not invent any information not found in the `germany_response`. Base your content strictly on the provided information. If any required information is missing from the source, simply omit that point rather than fabricating details.

            """},

            #  {
            #     "role": "user" , "content": f"""You have to write another paragraph about why Germany is the ideal destination to study {programme} based in which you will add 5-6 points explaining why you want to study in Germany and the programme {programme}, based on relevant examples from {germany_from_articles} """
            # },
            
            

            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},
            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},
            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},

            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": f"Be sure to replace the correct information from {res_text} including the name."},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is warm, simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on','delve', 'renowned', 'intricacies', 'close-knit', 'aligns', 'hands-on', 'enhance', 'foster', 'emphasis'"},

        ]

        )

        response_out = completion['choices'][0]['message']['content']
        st.write("**THE GENERATED SOP**" + '\n\n\n\n' + response_out)
        return response_out
    
    
    elif cgpa_score != 'unknown' and float(cgpa_score) < 7:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
            messages=[
             {"role": "user", "content": "Generate a statement of purpose (SOP) based on the provided details, following a specific structure and style."},

            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Here is the resume {res_text} from which you have to take information about the student and please make sure that you do not add things from your own and stick to this {res_text} througout the sop"},
            {"role": "user" , "content": f"Another important that you should keep in mind while making the sop is that the sop is a letter of motivation to study in this {university} , so you have to generate the sop which should contian strong motivational points for choosing {programme}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"University name: {university}"},

            {"role": "user", "content": f" For the Introduction's paragraph, you must be using content from the {res_text} and follow the same structure and catchy words from from the introduction paragraph of this template {template_text}"},


            {"role": "user", "content": f"""Fot the Academic and experience Paragraph too you will be using content from the {res_text}: Discuss the candidate's academic and professional background :
            IF CGPA score is more than 7 and his Bachelor graduation date is PAST : Mention his CGPA score.
            Mention all his acadamic education scores from the {res_text}
             Mention all the acadamic education scores as they are prensent in {res_text}
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list.
            You will be mentioning things like his certificates and Ilets score if they are present in the {res_text}
            """},


            {"role": "user", "content": f""" For the paragraph about the programme use strong motivational points for choosing {programme} and add information from the programme into this paragraph, 
                Show a STRONG motivation by evoking a storytelling from your past that led you want to study this programme.
                If the past education of yours is related to the programme you are applying then you can consider them adding in this paragraph too
                You have to choose MINIMUM two out of these points :
                Story from a family experience
                Story from a social phenomenon and you want to bring solutions
                Story from previous internship
                Story from the studied modules in your Academic cursus
                Story from one of your extra-curricular activites
                You will also be adding What are the benifits you will be getting from this {programme} and the {university} 
                You must also select some modules proposed by the programme and describe their contents, then explain why you are HIGHLY motivated to study them.
                Mention the learning outcomes from this {programme}, which include technical knowledge and interpersonal skills.     
                How this course will benefit me professionally and for persona growth as well.
                Explain Why this programme can help you build your career and prepare you for your future too.
                Please write this paragraph in the first-person tone where I express my strong motivation to study in Germany. 
                Start each sentence with 'I' or include 'my'
                **Important:** Please do not invent any information not found in the `programme_response`. Base your content strictly on the provided information. If any required information is missing from the source, simply omit that point rather than fabricating details.

            """},

            {"role": "user", "content": f"""For the paragraph about the University, you need to craft a compelling and personalized statement about why you want to apply to {university}. Write in the first person, as if you are the applicant. Incorporate detailed information from into this paragraph, addressing the following points:

            1. **University Ranking and Source:**
            - Include the exact ranking of {university} and cite the source of this ranking.

            2. **Student Body and Facilities:**
            - Mention the number of students at {university}.
            - Describe the facilities, faculty, and campus location in detail.

            3. **Personal Connections:**
            - Note if you have any relatives or friends at {university} and how that influences your choice.

            4. **Research Centers and Professors:**
            - Highlight any notable research centers and renowned professors associated with the program you are interested in.

            5. **City Attractions:**
            - Explain why you want to study in the city where {university} is located. Include some cool spots or attractions in the city that interest you.

            6. **Local Job Market:**
            - Provide examples of companies in the city that hire graduates from the {programme}.
            Based on the above 6 points, add the following points in the paragraph:
            7. Academic benifits:
            - Mention the academic benifits of studying in this {university}

            8. Personal benifits:
            - Mention the personal benifits of studying in this {university}

            9. Professional benifits:
            -Mention the professional benifits of studying in this {university}

            Please write this paragraph in the first-person tone where I express my strong motivation to study in Germany. 
            Start each sentence with 'I' or include 'my'

            Make sure to weave all these elements into a cohesive and engaging paragraph that reflects your genuine interest in both the university and the city. Be detailed, specific, and personal in your response.

            **Important:** Please do not invent any information not found in the `university_response`. Base your content strictly on the provided information. If any required information is missing from the source, simply omit that point rather than fabricating details.

            """},
            {"role": "user", "content": f"""For the paragraph which is about your future career perspectives and aspirations post-study, you should mention how the programme of {programme}
             will give you the opportunity to work in the futur job you want to do. You should mention names of professional positions linked to the programme you want to work in.
             You can also add that you aim to work in Germany"""},


            {"role": "user", "content": "Last Paragraph: A brief conclusion summarizing why you are the ideal applicant and show again your interest."},
            {"role": "user", "content": f"""
            For the paragraph about Germany, you must mention the following points and add information from into this paragraph :
            
            You strictly have to include at least 5-6 good points for this paragraph.
            Point 1. Explain why you chose to apply in Germany
            Point 2. Your Intention to stay in Germany because the field education comparative is far better to others destinations
            Point 3. Talk about the Good exposure, diversity and German culture 
            point 4. What are the reasons to choose Germany for masters/bachelors program?
            point 5. How studying in Germnay will help me grow professionally as well as personally
            point 6. Facts about Germany and its culture that led to the decision to study in Germany
            Please ensure each sentence transitions smoothly into the next, maintaining a logical flow throughout the document.
            Please write this paragraph in the first-person tone where I express my strong motivation to study in Germany. 
            Start each sentence with 'I' or include 'my'
            please can you write each sentence like "I want to …." instead of "Germany's education system is highly regarded for its quality and affordability" and please write the whole paragraph like this only.
            **Important:** Please do not invent any information not found in the `germany_response`. Base your content strictly on the provided information. If any required information is missing from the source, simply omit that point rather than fabricating details.

            """},

            #  {
            #     "role": "user" , "content": f"""You have to write another paragraph about why Germany is the ideal destination to study {programme} based in which you will add 5-6 points explaining why you want to study in Germany and the programme {programme}, based on relevant examples from {germany_from_articles} """
            # },
            
            

            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},
            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},
            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},

            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": f"Be sure to replace the correct information from {res_text} including the name."},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is warm, simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on','delve', 'renowned', 'intricacies', 'close-knit', 'aligns', 'hands-on', 'enhance', 'foster', 'emphasis'"},

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write("**THE GENERATED SOP**" + '\n\n\n\n' + response_out)
        return response_out
    else:
        completion = ai.ChatCompletion.create(
        #model="gpt-3.5-turbo-16k", 
        model = "gpt-4o-2024-05-13",
        temperature=ai_temp,
            messages=[
             {"role": "user", "content": "Generate a statement of purpose (SOP) based on the provided details, following a specific structure and style."},

            {"role": "user", "content": f"Template for guidance: {template_text}. You should reproduce the exact same structure"},
            {"role": "user", "content": f"Here is the resume {res_text} from which you have to take information about the student and please make sure that you do not add things from your own and stick to this {res_text} througout the sop"},
            {"role": "user" , "content": f"Another important that you should keep in mind while making the sop is that the sop is a letter of motivation to study in this {university} , so you have to generate the sop which should contian strong motivational points for choosing {programme}"},
            {"role": "user", "content": f"Study programme name: {programme}"},
            {"role": "user", "content": f"University name: {university}"},

            {"role": "user", "content": f" For the Introduction's paragraph, you must be using content from the {res_text} and follow the same structure and catchy words from from the introduction paragraph of this template {template_text}"},


            {"role": "user", "content": f"""Fot the Academic and experience Paragraph too you will be using content from the {res_text}: Discuss the candidate's academic and professional background :
            IF CGPA score is more than 7 and his Bachelor graduation date is PAST : Mention his CGPA score.
            Mention all his acadamic education scores from the {res_text}
             Mention all the acadamic education scores as they are prensent in {res_text}
            Mention ALL his work experiences and the different roles and responsibilities he held
            Retrieve ALL extra curricular activites and workshops he has attended and list them like a bulleted list.
            You will be mentioning things like his certificates and Ilets score if they are present in the {res_text}
            """},


            {"role": "user", "content": f""" For the paragraph about the programme use strong motivational points for choosing {programme} and add information from the programme into this paragraph, 
                Show a STRONG motivation by evoking a storytelling from your past that led you want to study this programme.
                If the past education of yours is related to the programme you are applying then you can consider them adding in this paragraph too
                You have to choose MINIMUM two out of these points :
                Story from a family experience
                Story from a social phenomenon and you want to bring solutions
                Story from previous internship
                Story from the studied modules in your Academic cursus
                Story from one of your extra-curricular activites
                You will also be adding What are the benifits you will be getting from this {programme} and the {university} 
                You must also select some modules proposed by the programme and describe their contents, then explain why you are HIGHLY motivated to study them.
                Mention the learning outcomes from this {programme}, which include technical knowledge and interpersonal skills.     
                How this course will benefit me professionally and for persona growth as well.
                Explain Why this programme can help you build your career and prepare you for your future too.
                Please write this paragraph in the first-person tone where I express my strong motivation to study in Germany. 
                Start each sentence with 'I' or include 'my'
                **Important:** Please do not invent any information not found in the `programme_response`. Base your content strictly on the provided information. If any required information is missing from the source, simply omit that point rather than fabricating details.

            """},

            {"role": "user", "content": f"""For the paragraph about the University, you need to craft a compelling and personalized statement about why you want to apply to {university}. Write in the first person, as if you are the applicant. Incorporate detailed information from into this paragraph, addressing the following points:

            1. **University Ranking and Source:**
            - Include the exact ranking of {university} and cite the source of this ranking.

            2. **Student Body and Facilities:**
            - Mention the number of students at {university}.
            - Describe the facilities, faculty, and campus location in detail.

            3. **Personal Connections:**
            - Note if you have any relatives or friends at {university} and how that influences your choice.

            4. **Research Centers and Professors:**
            - Highlight any notable research centers and renowned professors associated with the program you are interested in.

            5. **City Attractions:**
            - Explain why you want to study in the city where {university} is located. Include some cool spots or attractions in the city that interest you.

            6. **Local Job Market:**
            - Provide examples of companies in the city that hire graduates from the {programme}.
            Based on the above 6 points, add the following points in the paragraph:
            7. Academic benifits:
            - Mention the academic benifits of studying in this {university}

            8. Personal benifits:
            - Mention the personal benifits of studying in this {university}

            9. Professional benifits:
            -Mention the professional benifits of studying in this {university}

            Please write this paragraph in the first-person tone where I express my strong motivation to study in Germany. 
            Start each sentence with 'I' or include 'my'

            Make sure to weave all these elements into a cohesive and engaging paragraph that reflects your genuine interest in both the university and the city. Be detailed, specific, and personal in your response.

            **Important:** Please do not invent any information not found in the `university_response`. Base your content strictly on the provided information. If any required information is missing from the source, simply omit that point rather than fabricating details.

            """},
            {"role": "user", "content": f"""For the paragraph which is about your future career perspectives and aspirations post-study, you should mention how the programme of {programme}
             will give you the opportunity to work in the futur job you want to do. You should mention names of professional positions linked to the programme you want to work in.
             You can also add that you aim to work in Germany"""},


            {"role": "user", "content": "Last Paragraph: A brief conclusion summarizing why you are the ideal applicant and show again your interest."},
            {"role": "user", "content": f"""
            For the paragraph about Germany, you must mention the following points and add information from the news into this paragraph :
            
            You strictly have to include at least 5-6 good points for this paragraph.
            Point 1. Explain why you chose to apply in Germany
            Point 2. Your Intention to stay in Germany because the field education comparative is far better to others destinations
            Point 3. Talk about the Good exposure, diversity and German culture 
            point 4. What are the reasons to choose Germany for masters/bachelors program?
            point 5. How studying in Germnay will help me grow professionally as well as personally
            point 6. Facts about Germany and its culture that led to the decision to study in Germany
            Please ensure each sentence transitions smoothly into the next, maintaining a logical flow throughout the document.
            Please write this paragraph in the first-person tone where I express my strong motivation to study in Germany. 
            Start each sentence with 'I' or include 'my'
            please can you write each sentence like "I want to …." instead of "Germany's education system is highly regarded for its quality and affordability" and please write the whole paragraph like this only.
            **Important:** Please do not invent any information not found in the `germany_response`. Base your content strictly on the provided information. If any required information is missing from the source, simply omit that point rather than fabricating details.

            """},

            #  {
            #     "role": "user" , "content": f"""You have to write another paragraph about why Germany is the ideal destination to study {programme} based in which you will add 5-6 points explaining why you want to study in Germany and the programme {programme}, based on relevant examples from {germany_from_articles} """
            # },
            
            

            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},
            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},
            { "role": "user" , "content": f"Don't consider the above instructions as the order of the paragraphs of the statement of purpose but use {template_text} for the  structure and order instead"},

            {"role": "user", "content": "Finish with a closing line with consideration of this statement of purpose and add your name signature bellow"},
            {"role": "user", "content": f"Be sure to replace the correct information from {res_text} including the name."},
            {"role": "user", "content": "Please ensure each paragraph transitions smoothly into the next, maintaining a logical flow throughout the document."},
            {"role": "user", "content": "The statement of purpose should consist of seven paragraphs, totaling a minimum of 500 words, using simple language that appears human-written."},
            {"role": "user", "content": "MOST IMPORTANT : Make sure the tone is warm, simple and human-like. Don't use the following words : 'cutting-edge', 'leverage', 'honed/hone', 'appealing', 'hands-on','delve', 'renowned', 'intricacies', 'close-knit', 'aligns', 'hands-on', 'enhance', 'foster', 'emphasis'"},

        ]
        )

        response_out = completion['choices'][0]['message']['content']
        st.write("**THE GENERATED SOP**" + '\n\n\n\n' + response_out)
        return response_out








# Collect data from each fragment
async def collect_data_fragment_diffbot(base_url, fragment_links):
    all_data = {}
    for fragment in fragment_links:
        full_url = f'{base_url}{fragment}'
        if full_url:
            content = await get_url_content(full_url)
            if content:
                all_data[fragment] = content
        else:
            print(f"Nothing found for {full_url}")
            return None
    return all_data

# Main function
def return_data_withfragments(query):
    if extract_2_urls(query) is not None:
        url1, url2 = extract_2_urls(query)
    # Base URL
        bases_urls =  [url1, url2]
        programme_content = []
        for base_url in bases_urls:
            fragment_links = get_fragment_links(base_url)
            if fragment_links:
                all_data = collect_data(base_url, fragment_links)                                       
                for fragment, data in all_data.items():
                    print(f'Data from {fragment}:')
                    print(data)
                    programme_content.append(data)
            else:
                print('No fragment links found.')
        return ("\n".join(programme_content))
    else:
        return None
     
# Main function
async def return_data3(query):  
    url1, url2, url3 = await extract3urls(query)  # Await extract3urls here
    
    # Create tasks for fetching content from URLs asynchronously
    if all([url1, url2, url3]):
        task1 = asyncio.create_task(get_url_content(url1))
        task2 = asyncio.create_task(get_url_content(url2))
        task3 = asyncio.create_task(get_url_content(url3))
        
        # Wait for all tasks to complete
        content1, content2, content3 = await asyncio.gather(task1, task2, task3)
            
        # Filter out any None values
        tablo = [content for content in [content1, content2, content3] if content is not None]
        
        if tablo:
            result = '\n\n\n'.join(tablo)
            # print(result)
            # Or return the result if this is inside a function
            return result
    else:
        print("One or more URLs are None")
        return None
        

async def return_data3_orderofresults(query):  
    url1, url2, url3 = await extract3urls_orderofresults(query)  # Await extract3urls here
    
    # Create tasks for fetching content from URLs asynchronously
    if all([url1, url2, url3]):
        task1 = asyncio.create_task(get_url_content(url1))
        task2 = asyncio.create_task(get_url_content(url2))
        task3 = asyncio.create_task(get_url_content(url3))
        
        # Wait for all tasks to complete
        content1, content2, content3 = await asyncio.gather(task1, task2, task3)
            
        # Filter out any None values
        tablo = [content for content in [content1, content2, content3] if content is not None]
        
        if tablo:
            result = '\n\n\n'.join(tablo)
            # print(result)
            # Or return the result if this is inside a function
            return result
    else:
        print("One or more URLs are None")
        return None    

# Function to load the scraped data from the pickle file
async def load_scraped_data():
    if os.path.exists(PICKLE_FILE):
        with open(PICKLE_FILE, 'rb') as file:
            return pickle.load(file)
    return {}

# Function to save the scraped data to the pickle file
async def save_scraped_data(data):
    with open(PICKLE_FILE, 'wb') as file:
        pickle.dump(data, file)


async def extract_first_url_1(query):
    url = f'https://www.google.com/search?q={requests.utils.quote(query)}'
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }

    async with aiohttp.ClientSession() as session:
        async with session.get(url, headers=headers) as response:
            if response.status == 200:
                html = await response.text()
                soup = BeautifulSoup(html, 'html.parser')
                result = soup.find('div', {'class': 'yuRUbf'})
                if result:
                    print(f"The Unique link for {query} :", result.find('a')['href'])
                    return result.find('a')['href']
                else:
                    print("No results found.")
                    return None
            else:
                print("Failed to retrieve the website content.")
                return None

    
# Function to scrape content from a URL using the Diffbot API
async def scrape_url(url):
    api_url = f"{DIFFBOT_API_URL}?token={DIFFBOT_API_TOKEN}&url={url}"
    response = requests.get(api_url)
    
    if response.status_code == 200:
        response_data = response.json()
        if 'objects' in response_data and len(response_data['objects']) > 0:
            html_content = response_data['objects'][0].get('html', 'HTML content not available')
            if html_content == 'HTML content not available':
                print(f" for {url} : HTML content not available in the response.")
                return None
            return html_content
        else:
            print(f" for {url} : HTML content not available in the response.")
            return None
    else:
        print(f"Failed to retrieve data from {url}: {response.status_code}")
        return None

# Function to clean HTML content and extract text
async def clean_html(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    return soup.get_text(separator='\n', strip=True)

# Function to get content from a URL, using cached data if available
async def get_url_content(url):
    scraped_data = await load_scraped_data()
    
    if url in scraped_data:
        print(f"Using cached data for {url}")
        return scraped_data[url]
    else:
        content = await scrape_url(url)
        if content:
            text_content = await clean_html(content)
            if text_content:  # Check if the content is not empty
                scraped_data[url] = text_content
                await save_scraped_data(scraped_data)
            return text_content
        return None


async def return_data1(query):
    url1 = await extract_first_url_1(query)
    print(f"for query: {query}  :", url1)
    if url1:
        content = await get_url_content(url1)
        return content
    else:
        print("Nothing found for url1")
        return None

async def return_data1_with_fragments(query):
    url1 = await extract_first_url_1(query)
    print(f"for query: {query}  :", url1)
    programme_content = []
    if url1:
        fragment_links = get_fragment_links(url1)
        if fragment_links:
            all_data = await collect_data_fragment_diffbot(url1, fragment_links)
            if all_data:                                    
                for fragment, data in all_data.items():
                    print(f'Data from {fragment}:')
                    print(data)
                    programme_content.append(data)
                else:
                    print('No content found.')
        
                return ''.join(programme_content)
            else:
                print('No data found.')
                return None
        else:
            print('No fragment found')
            return None
    else:
        print(f"No url found for the query {query}")
        return None

async def return_data1_see(query):
    url1 = await extract_1_url_from_urls(query)
    print(f"for query: {query}  :", url1)
    if url1:
        content = await get_url_content(url1)
        return content
    else:
        print("Nothing found for url1")
        return None


async def retrieve_cgpa_score(content_resume):
    
    # Check if the entity_ruler already exists in the pipeline
    if 'entity_ruler' in nlp.pipe_names:
        ruler = nlp.get_pipe('entity_ruler')
    else:
        ruler = nlp.add_pipe("entity_ruler")

    patterns = [
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}]},  # Matches "9.6 CGPA"
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"IS_PUNCT": True, "OP": "?"}]},  # Matches "CGPA 9.6" and "CGPA 9.5" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}]},  # Matches "CGPA - 9.26" or "CGPA 9.5" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}]},  # Matches "9.5 CGPA" or "9.5 CGPA -" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"IS_PUNCT": True, "OP": "?"}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}]},  # Matches "9.5 - CGPA" or "9.5 . CGPA" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}, {"LOWER": {"IN": ["-", "–"]}}, {"IS_SPACE": True, "OP": "?"}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}]}, # Matches "CGPA - 9.26" or "CGPA -9.5" (with optional punctuation and space)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\w*\\d+(\\.\\d+)?"}}, {"LOWER": {"REGEX": "(cgpa|sgpi\\w*)"}}]},
        {"label": "CGPA", "pattern": [{"LOWER": {"REGEX": "(\\w*cgpa|sgpi)"}}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?\w*"}}]}
    ]

    ruler.add_patterns(patterns)

    lines = content_resume.split('\n')
    score = "unknown"
    for index, line in enumerate(lines):
        doc = nlp(line)
        found_cgpa = False
        for ent in doc.ents:
            if ent.label_ == 'CGPA':
                found_cgpa = True
                print(ent.text)
                numerical_score = nlp_e(ent.text)
                for token in numerical_score:
                    if token.ent_type_ == 'CARDINAL':
                        score = token.text
                break  # Stop looping over entities once CGPA label is found
        if found_cgpa:
            break  # Stop looping over lines once CGPA label is found

    return score


def retrieve_cgpa_score_sync(content_resume):
    
    # Check if the entity_ruler already exists in the pipeline
    if 'entity_ruler' in nlp.pipe_names:
        ruler = nlp.get_pipe('entity_ruler')
    else:
        ruler = nlp.add_pipe("entity_ruler")

    patterns = [
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}]},  # Matches "9.6 CGPA"
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"IS_PUNCT": True, "OP": "?"}]},  # Matches "CGPA 9.6" and "CGPA 9.5" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}]},  # Matches "CGPA - 9.26" or "CGPA 9.5" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}]},  # Matches "9.5 CGPA" or "9.5 CGPA -" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}, {"IS_PUNCT": True, "OP": "?"}, {"LOWER": {"IN": ["cgpa", "sgpi"]}}]},  # Matches "9.5 - CGPA" or "9.5 . CGPA" (with optional punctuation)
        {"label": "CGPA", "pattern": [{"LOWER": {"IN": ["cgpa", "sgpi"]}}, {"IS_PUNCT": True, "OP": "?"}, {"LOWER": {"IN": ["-", "–"]}}, {"IS_SPACE": True, "OP": "?"}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?"}}]}, # Matches "CGPA - 9.26" or "CGPA -9.5" (with optional punctuation and space)
        {"label": "CGPA", "pattern": [{"TEXT": {"REGEX": "\w*\\d+(\\.\\d+)?"}}, {"LOWER": {"REGEX": "(cgpa|sgpi\\w*)"}}]},
        {"label": "CGPA", "pattern": [{"LOWER": {"REGEX": "(\\w*cgpa|sgpi)"}}, {"TEXT": {"REGEX": "\\d+(\\.\\d+)?\w*"}}]}
    ]

    ruler.add_patterns(patterns)

    lines = content_resume.split('\n')
    score = "unknown"
    for index, line in enumerate(lines):
        doc = nlp(line)
        found_cgpa = False
        for ent in doc.ents:
            if ent.label_ == 'CGPA':
                found_cgpa = True
                print(ent.text)
                numerical_score = nlp_e(ent.text)
                for token in numerical_score:
                    if token.ent_type_ == 'CARDINAL':
                        score = token.text
                break  # Stop looping over entities once CGPA label is found
        if found_cgpa:
            break  # Stop looping over lines once CGPA label is found

    return score

    
def create_word_document(phrase, font_name, font_size):  
    doc = Document()
    doc.add_paragraph(phrase)

    # Set font properties
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            
        # Set paragraph alignment to justified
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    return doc


def save_doc_to_buffer(doc):
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
    
    
def read_pdf(file):
    pdf_reader = PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


def generate_random_templates(directory_path):
    # List all files in the directory
    files = os.listdir(directory_path)
    # Filter out the PDF files
    pdf_files = [file for file in files if file.endswith('.pdf')]
    
    # Check if there are any PDF files
    if not pdf_files:
        print("No PDF files found in the directory.")
        return
    
    # Randomly select a PDF file
    selected_pdf = random.choice(pdf_files)
    print(f"Selected PDF: {selected_pdf}")
    
    # Full path to the selected PDF file
    pdf_path = os.path.join(directory_path, selected_pdf)
    return read_pdf(pdf_path)
    

def read_docx(file):
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


async def return_content_germany_articles(data, element):
    text_result = []
    for item in data.get('data', []):
        entity = item.get('entity', {})
        title = entity.get('title', 'No title')
        text = entity.get('text', 'No text')

        if len(title) < 40 and not any(ele.lower() in title.lower() for ele in element):
            text_result.append(f"Title: {title}\nText: {text}\n")
    
    # If there are suitable articles, randomly select one
    if text_result:
        return random.choice(text_result)
    
    # If no suitable article is found, return a default message
    return "No suitable article found."


async def extract_program_name(text):
    doc = nlp(text)
    program_name = None
    start_index = 0

    # List of patterns to check at the start of the string
    patterns = [
        "master in ",
        "masters in ",
        "master of ",
        "masters of ",
        "master of science in ",
        "master of science of ",
        "masters of science in ",
        "masters of science of ",
        "masters of science",
        "master of science",
        "msc",
        "msc of",
        "msc in",
        "master",
        "masters",
        "m.sc",
        "m.sc."
    ]

    # Convert the first few tokens to lowercase and join them to form the beginning of the string
    beginning_of_string = " ".join([token.text.lower() for token in doc[:4]])

    # Check if the beginning of the string matches any of the specified patterns
    if not any(beginning_of_string.startswith(pattern) for pattern in patterns):
        return text  # If no pattern matches, return the entire string
    
    # Iterate through tokens to find the appropriate pattern
    for i, token in enumerate(doc):

        # Case 1: "Masters in X" or "Master in X"
        if token.text.lower() in ["masters", "master"] and i + 2 < len(doc) or "msc":
            if doc[i+1].text.lower() == "in":
                start_index = i + 2
                break
        
        # Case 2: "Masters of X" or "Master of X"
        if token.text.lower() in ["masters", "master"] and i + 2 < len(doc) or "msc":
            if doc[i+1].text.lower() == "of":
                # Check if "science in" follows, "science of" or just "science"
                if doc[i+2].text.lower() == "science":
                    if i + 3 < len(doc) and doc[i+3].text.lower() == "in":
                        start_index = i + 4  # "Master of science in X"
                    if i + 3 < len(doc) and doc[i+3].text.lower() == "of":
                        start_index = i + 4  # "Master of science in X"
                    
                    else:
                        start_index = i + 3  # "Master of science X"
                else:
                    start_index = i + 2  # "Master of X"
                break
        
        # Case 3: "Masters X" or "Master X"
        if token.text.lower() in ["masters", "master"] or "msc":
            start_index = i + 1
            break

    # Extract the program name from the identified index
    if start_index > 0:
        program_name = " ".join([t.text for t in doc[start_index:]])
    
    # Additional functionality to compare the number of words and characters around "and"
    if program_name and "and" in program_name.lower():
        parts = program_name.lower().split("and")
        left_part = parts[0].strip()
        right_part = parts[1].strip()

        left_word_count = len(left_part.split())
        right_word_count = len(right_part.split())

        if right_word_count > left_word_count:
            program_name = right_part
        elif left_word_count > right_word_count:
            program_name = left_part
        else:
            # Compare the number of characters
            left_char_count = len(left_part)
            right_char_count = len(right_part)

            if right_char_count > left_char_count:
                program_name = right_part
            elif left_char_count > right_char_count:
                program_name = left_part
            else:
                # Return either side randomly if both word and character counts are equal
                program_name = random.choice([left_part, right_part])

    return program_name





#& below : every settings linked to log in and cookies : 

#* Turn Streamlit's st.cache deprecation warnings into exceptions
warnings.filterwarnings("error", message=".*st.cache is deprecated.*")

#* Initialize the cookie manager
cookies = EncryptedCookieManager(password="your-secure-password")

#* Ensure cookies are ready before proceeding
if not cookies.ready():
    st.stop()  # Stop the script until the cookies are ready

#* Check if the user is already logged in via cookies
if "logged_in" not in st.session_state:
    logged_in_cookie = cookies.get("logged_in")
    st.session_state.logged_in = logged_in_cookie == "true"
    st.session_state.user_email = cookies.get("user_email") if st.session_state.logged_in else None


if not st.session_state.logged_in:
    st.image(logo_path, width=150)
    st.title("Welcome to YES WIZARD")
    st.subheader("Login")
    email = st.text_input('Please enter your email address')
    password = st.text_input('Please enter your password' , type= 'password')

    if st.button('Login'):
        try:
            user = auth.sign_in_with_email_and_password(email , password)
            st.session_state.logged_in = True
            st.session_state.user = user['email']

            # Set cookies to persist the login state for 24 hours
            expires_at = (datetime.now() + timedelta(seconds=60)).isoformat()
            cookies["logged_in"] = "true"
            cookies["user_email"] = user['email']
            cookies["expires_at"] = expires_at
            cookies.save()

            st.success('Welcome!')
            st.rerun()  # Refresh the page to show the main content
        except Exception as e:
            print(f"Error: {e}")
            st.error("Wrong Credentials!")
    
else:
    
    col1, col2, col3 = st.columns([1, 6, 1])
    
    with col3:
        # Logout button
        if st.button('Logout'):
                # Clear session state
                st.session_state.logged_in = False
                st.session_state.user_email = None

                # Delete the cookies to log out
                del cookies["logged_in"]
                del cookies["user_email"]
                del cookies["expires_at"]
                cookies.save()

                st.rerun()  # Refresh the page to show the login screen
    
    # Display the logo
    st.image(logo_path, width=150)  # Adjust the width as needed
    st.markdown(f"""
    # YES WIZARD

    """
    )

    # radio for upload or copy paste option         

    res_file = st.file_uploader('📁 Upload your resume in pdf or docx format')
        
    if res_file is not None:
        if res_file.name.endswith('.pdf'):
            pdf_reader = PdfReader(res_file)

            # Collect text from pdf
            res_text = ""
            for page in pdf_reader.pages:
                res_text += page.extract_text()
        
        elif res_file.name.endswith('.docx'):
            doc_reader = Document(res_file)

            res_text = ""
            for para in doc_reader.paragraphs:
                res_text += para.text + " "


                for table in doc_reader.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            res_text += cell.text + " "
                        res_text += "\n"

        else:
            st.error("Unsupported file format. Please upload a PDF or DOCX file.")
        
    

                
                

    with st.form('input_form'):
        # other inputs
        programme = st.text_input('Programme/Course name')
        university = st.text_input('University name')
        ai_temp = st.number_input('AI Temperature (0.0-1.0) Input how creative the API can be',value=.6)
        
        
        
        
        # # submit button
        submitted = st.form_submit_button("Generate the SOP")

    # if the form is submitted run the openai completion   
    if submitted:    
        async def main():

            research = university + " " + programme
            extract_programme_name = await extract_program_name(programme)
            print(extract_programme_name)


        #! DQL OPERATIONS TO RETRIVE CONTENT OF ARTICLES ABOUT GERMANY AND THE PROGRAMME 
            # Base URL
            base_url = "https://kg.diffbot.com/kg/v3/dql"

            # DQL query
            dql_query = f"type:Article title:'Germany' sentiment>=0.80 title:\"{extract_programme_name}\""
            # Encode the DQL query to be URL-safe
            encoded_query = urllib.parse.quote(dql_query)
            url = f"{base_url}?token={DIFFBOT_API_TOKEN}&query={encoded_query}"

            headers = {"Accept": "application/json"}

            element = ['Top', 'Best']
            
            async with aiohttp.ClientSession() as session:
                async with session.get(url, headers=headers) as response:
                    if response is not None:
                        data = await response.json()
                        germany_from_articles = await return_content_germany_articles(data, element)
                        print(f"germany_from_articles: {germany_from_articles}")



            response_sop = await generate_sop_version1( generate_random_templates('templates4'), res_text,programme,university)

            
            # Combine responses
            response = ''.join(response_sop) 
            final_sop_and_responses = response + response_sop 

            
            doc_download1 = create_word_document(final_sop_and_responses, 'Arial', 11)
            st.download_button(
                label="Download SOP",
                data=save_doc_to_buffer(doc_download1),
                file_name=f"{res_file.name}_{university}_{programme}_SOP.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        asyncio.run(main())